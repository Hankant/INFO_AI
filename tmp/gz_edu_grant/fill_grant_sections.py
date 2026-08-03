from pathlib import Path
import shutil

from docx import Document


SRC = Path(r"E:\Info_AI\tmp\gz_edu_grant\template.docx")
OUT = Path(r"E:\Info_AI\tmp\gz_edu_grant\grant_sections_draft.docx")


TITLE_CN = "生成式人工智能重构高风险决策的信息校准与依赖学习机制研究"
TITLE_EN = "Generative AI, Information Calibration, and Reliance Learning in High-Risk Decision-Making"
RESEARCH_DIRECTION = "生成式人工智能与高风险决策；信息校准；人机协作依赖学习"
ABSTRACT = (
    "生成式人工智能正从信息检索工具转向判断与辅助决策基础设施。本项目围绕“AI如何重构高风险决策”的关键问题，"
    "聚焦两个相互衔接的机制：一是 answer-first summary 及其来源标签如何改变用户对额外证据价值、信息充分性和后续信念更新的判断；"
    "二是在重复反馈中，早期AI错误是否会引发相对于人类来源更强、更持久的低依赖与低授权。项目将采用 same-summary "
    "source-label 实验、重复预测/判断实验、贝叶斯基准以及简化 EWA/来源特异性学习模型，识别生成式AI对信息校准与依赖学习的影响机制，"
    "并为高校及公共部门在AI辅助判断场景中的界面设计、使用规范和风险治理提供经验证据。"
)
KEYWORDS = "生成式人工智能, 信息校准, 依赖学习"

JUSTIFICATION = """生成式人工智能正在显著改变人们获取信息、整合证据和做出判断的方式。与传统搜索或一般决策支持系统不同，生成式AI往往以 answer-first 的方式直接给出总结、建议或默认判断，这意味着用户在尚未逐条接触原始证据之前，已经面对一个被系统压缩、组织和叙事化的结果。在高风险决策场景中，这一变化带来的问题不再只是“AI是否更准确”，而是“AI如何改变人对证据价值、信息充分性和后续依赖边界的判断”。因此，围绕生成式AI展开研究，必须把分析单位从单次正确率比较推进到信息校准与依赖学习过程。

现有 trust in automation、algorithm aversion、algorithm appreciation 与 AI-assisted decision-making 文献，为理解人类如何评价算法系统提供了重要基础，但大多集中在单轮建议采纳、局部置信校准或短期错误反应层面。第一，这些研究较少区分“搜索减少”究竟意味着理性的信息压缩，还是有害的信息获取偏离。第二，它们较少考察人在重复反馈中如何学习AI应当扮演什么角色，以及早期AI错误是否会被赋予比人类来源错误更高的负向学习权重。第三，现有研究往往把信息获取、后验更新和依赖行为分别处理，尚未把 evidence-to-belief 与 belief-to-reliance 两个关键环节置于同一研究计划中加以识别。

本项目因此聚焦两个前后衔接的问题。其一，在初始 summary 内容相同、额外搜索成本相同的条件下，AI 来源标签是否会改变用户对额外证据价值、主观信息获得感、epistemic ownership 和后续信念更新的判断。其二，在重复反馈任务中，即使 AI 与 human source 的总体表现等价，早期AI错误是否仍会引发更强、更持久的后续低依赖。前者对应高风险决策中的信息校准问题，后者对应动态依赖校准问题。将两者放在同一项目中研究，可以更系统地解释生成式AI如何重构高风险决策中的“证据—信念—依赖”链条。

从现实意义看，生成式AI已经进入高校科研、政策分析、内容审核、风险筛查、例外复核与专业判断等多种场景。无论在高校科研训练还是公共部门辅助决策中，仅凭平均准确率并不足以判断AI系统是否被恰当使用。真正关键的是：用户是否过早停止搜索、是否误判自身信息掌握程度、是否在少量早期错误后长期回到过强的人为保守状态。本项目有助于为高校和公共部门建立更审慎的AI使用规范、界面设计原则与依赖审计框架提供实验依据，也可为广州在智能治理与高层次人才培养中的制度建设提供可操作的知识支持。

主要参考文献：
[1] Lee J D, See K A. Trust in Automation: Designing for Appropriate Reliance[J]. Human Factors, 2004, 46(1): 50-80.
[2] Dietvorst B J, Simmons J P, Massey C. Algorithm Aversion: People Erroneously Avoid Algorithms after Seeing Them Err[J]. Journal of Experimental Psychology: General, 2015, 144(1): 114-126.
[3] Logg J M, Minson J A, Moore D A. Algorithm Appreciation: People Prefer Algorithmic to Human Judgment[J]. Organizational Behavior and Human Decision Processes, 2019, 151: 90-103.
[4] Parasuraman R, Sheridan T B, Wickens C D. A Model for Types and Levels of Human Interaction with Automation[J]. IEEE Transactions on Systems, Man, and Cybernetics-Part A, 2000, 30(3): 286-297.
[5] Zhang Y, Liao Q V, Bellamy R K E. Effect of Confidence and Explanation on Accuracy and Trust Calibration in AI-Assisted Decision Making[C]//FAccT 2020: 295-305.
[6] Ma S, Lei Y, Wang X, et al. Who Should I Trust: AI or Myself?[C]//CHI 2023: Article 759.
[7] Camerer C, Ho T H. Experience-Weighted Attraction Learning in Normal Form Games[J]. Econometrica, 1999, 67(4): 827-874.
[8] Mozannar H, Sontag D. Consistent Estimators for Learning to Defer to an Expert[C]//ICML 2020: 7076-7087.
[9] Arthur W B. Competing Technologies, Increasing Returns, and Lock-in by Historical Events[J]. The Economic Journal, 1989, 99(394): 116-131.
[10] Vaccaro M, Almaatouq A, Malone T W. When Combinations of Humans and AI Are Useful[J]. Nature Human Behaviour, 2024, 8: 2293-2303."""

RESEARCH_PLAN = """（一）研究内容、研究目标以及拟解决的关键问题。
本项目围绕“生成式人工智能如何重构高风险决策中的信息校准与依赖学习”展开，拟设置两个彼此衔接的实证研究模块。第一模块聚焦信息校准，在初始 summary 内容相同、后续证据池相同、额外搜索成本相同的条件下，仅操纵是否标注为“AI生成”，识别AI来源标签是否改变用户对额外证据价值、信息充分性、epistemic ownership 和后续信念更新的判断。第二模块聚焦动态依赖学习，在重复反馈任务中比较 AI early error、AI late error、AI distributed error 与 human-source error 等条件，在总体表现等价的前提下识别早期AI错误是否会引发更强、更持久的后续低依赖。项目总体目标是构建 evidence-to-belief 与 belief-to-reliance 相衔接的研究框架，并提出面向高校和公共部门AI辅助判断场景的设计与治理建议。拟解决的关键问题包括：AI标签是否在不改变信息内容的情况下改变用户对继续搜索必要性的判断；AI导致的搜索减少究竟是理性信息压缩还是信息获取偏离；早期AI错误是否比人类来源错误更易获得过强的负向学习权重；如何区分路径依赖、一般性惯性和先验偏见对后续依赖的影响。

（二）本项目的特色与创新之处。
本项目的创新主要体现在三个方面。第一，研究对象从一般性的“AI信任”推进到“信息校准与依赖学习”，不再只问用户是否接受某次建议，而是追问AI如何改变证据获取阈值、后验更新和后续依赖轨迹。第二，采用 same-summary source-label design，在 summary 内容不变、边际搜索成本相同的条件下识别 AI 标签效应，从方法上更清楚地区分内容效应、来源效应和成本效应。第三，将单轮贝叶斯基准实验与多轮反馈实验结合起来，既能识别AI是否诱发信息获取偏离，也能识别早期AI错误是否在后续多轮中造成持续性低依赖，从而把信息搜索、行为学习与人机协作边界置于同一框架中讨论。

（三）采用的研究方法。
本项目采用行为实验、规范性基准比较与动态学习模型相结合的方法。第一模块拟开展被试间实验，设置 no-summary control、unlabeled summary 与 AI-labeled summary 等条件，并区分 representative summary 与 incomplete summary。通过固定证据库、预设 likelihood ratio 和 log-odds 形式的贝叶斯基准，分别计算 conditional Bayesian deviation 与 full-information deviation，用以衡量用户对已接触证据的吸收质量，以及其相对于完整证据集的偏离程度。与此同时，测量 confidence、perceived information acquisition、epistemic ownership 和 perceived sufficiency of evidence 等变量，以识别主观—客观校准错配。第二模块拟采用有限重复窗口下的重复预测/判断实验。参与者在多轮任务中反复面对 AI 与 human source 的建议或判断结果，并根据每轮反馈更新后续依赖。研究通过在总体表现等价的前提下操控 early/late/distributed error 的来源与时点，结合 reduced-form 动态分析和简化 EWA / source-specific learning 模型，估计 AI error 与 human-source error 的负向更新权重是否存在系统性差异。

（四）技术路线以及项目效益分析。
项目技术路线分为“理论梳理—预实验校准—正式实验—综合整合”四步。首先围绕 information calibration 与 dynamic reliance learning 两个核心问题完成文献整合与研究设计收束。其次开展预实验，对 summary 材料、证据库、风险判断任务、重复反馈任务及关键测量指标进行校准。再次分别完成信息校准实验与动态依赖实验，形成两个模块的核心数据。最后通过统一框架整合 evidence-to-belief 与 belief-to-reliance 结果，提炼面向高校和公共部门的AI辅助判断设计建议。项目效益主要体现为：为高校科研训练、研究方法教学和公共部门辅助判断提供关于AI呈现方式、错误恢复机制和依赖审计的机制证据。

（五）项目预期风险及规避措施。
本项目的主要风险包括：AI标签效应可能较弱、多轮任务可能出现理解偏差或疲劳、动态学习模型可能存在参数可识别性不足，以及正式实验样本量不足等。为此，项目将通过预实验优化 summary 文本、标签措辞和任务情境，设置练习轮和理解检验，优先将 reduced-form 结果作为主证据，并在结构估计前进行参数恢复检验；同时采用“预实验—正式实验”两阶段方案，根据效应大小动态调整样本量。

（六）年度研究计划及预期研究结果。
第一阶段完成文献综述、理论框架和研究设计定稿，搭建固定证据库与 summary 材料，完成信息校准模块预实验和任务修订。第二阶段完成信息校准模块正式实验和初步论文写作，并开展重复反馈任务预实验，校准 error timing、performance parity 与测量设计，参加至少1次国内学术会议交流。第三阶段完成动态依赖学习模块正式实验与综合分析，形成完整研究报告、论文初稿和治理建议稿，再参加至少1次学术交流活动。预期研究结果包括：形成围绕“信息校准—依赖学习”主线的阶段性论文与研究报告，构建可复用的实验材料和分析框架，并提出关于AI辅助判断场景中证据呈现、错误恢复和依赖审计的政策建议。

（七）项目预期研究成果转移转化情况。
本项目虽属于人文社会科学研究，但具有明确的应用转化潜力。项目可为高校科研训练、研究方法课程、AI素养教育和公共部门辅助判断场景提供可操作的使用原则，包括：何种界面设计会降低用户继续搜索的意愿、何种错误反馈机制更有利于避免长期低依赖、何种指标可用于审计AI系统是否被恰当依赖。项目成果可转化为研究报告、培训材料、内部讲座与规范建议，为广州高校和公共部门开展AI工具治理提供经验支持。"""


def main() -> None:
    shutil.copyfile(SRC, OUT)
    doc = Document(str(OUT))

    doc.paragraphs[11].text = f"项目名称：{TITLE_CN}"

    overview = doc.tables[6]
    overview.cell(0, 4).text = TITLE_CN
    overview.cell(1, 4).text = TITLE_EN
    overview.cell(3, 4).text = RESEARCH_DIRECTION
    overview.cell(4, 6).text = "3年"
    overview.cell(5, 2).text = ABSTRACT
    overview.cell(6, 3).text = KEYWORDS

    doc.tables[7].cell(0, 0).text = JUSTIFICATION
    doc.tables[8].cell(0, 0).text = RESEARCH_PLAN

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
